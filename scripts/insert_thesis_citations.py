from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


DOC_PATH = Path(
    "output/doc/本科毕业论文初稿_20260409/附件6：大连海事大学本科毕业论文模板（理工类）-初稿.docx"
)

CITATIONS: dict[int, str] = {
    40: "[1]",
    44: "[2]",
    45: "[2,3]",
    53: "[4-9]",
    56: "[10]",
    59: "[2,3]",
    86: "[9,10]",
    96: "[11]",
    98: "[7-9]",
    104: "[4-8,10]",
    113: "[9,10]",
    116: "[10,11]",
    123: "[10]",
    126: "[12]",
    127: "[13]",
}


def set_run_times_new_roman(run) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    run.font.superscript = True


def replace_paragraph_with_citation(paragraph, citation: str) -> None:
    base_text = paragraph.text
    base_text = re.sub(r"\[[0-9,\-]+\]\s*$", "", base_text).rstrip()
    paragraph.clear()
    paragraph.add_run(base_text)
    citation_run = paragraph.add_run(citation)
    set_run_times_new_roman(citation_run)


def main() -> None:
    document = Document(DOC_PATH)
    for index, citation in CITATIONS.items():
        replace_paragraph_with_citation(document.paragraphs[index], citation)
    temp_path = DOC_PATH.with_name(f"{DOC_PATH.stem}.cite.tmp.docx")
    if temp_path.exists():
        temp_path.unlink()
    document.save(temp_path)
    temp_path.replace(DOC_PATH)


if __name__ == "__main__":
    main()

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


DOC_PATH = Path(
    "output/doc/本科毕业论文初稿_20260409/附件6：大连海事大学本科毕业论文模板（理工类）-初稿.docx"
)

REFERENCES = [
    "[1] Vaswani A, Shazeer N, Parmar N, et al. Attention Is All You Need[C]. Advances in Neural Information Processing Systems, 2017",
    "[2] Lewis P, Perez E, Piktus A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks[C]. Advances in Neural Information Processing Systems, 2020",
    "[3] Robertson S, Zaragoza H. The Probabilistic Relevance Framework: BM25 and Beyond[J]. Foundations and Trends in Information Retrieval, 2009, 3(4):333-389",
    "[4] Vercel. Next.js Documentation[OL]. https://nextjs.org/docs, 2026.04.10",
    "[5] Meta Platforms, Inc. React Documentation[OL]. https://react.dev/, 2026.04.10",
    "[6] OpenJS Foundation. Express Documentation[OL]. https://expressjs.com/, 2026.04.10",
    "[7] Fette I; Melnikov A. The WebSocket Protocol: RFC 6455[EB/OL]. https://www.rfc-editor.org/rfc/rfc6455, 2011.12",
    "[8] FastAPI Documentation[OL]. https://fastapi.tiangolo.com/, 2026.04.10",
    "[9] RAG_GPTV1.0 项目多服务开发文档[Z]. 项目技术文档, 2026",
    "[10] RAG_GPTV1.0 项目说明文档[Z]. 项目技术文档, 2026",
    "[11] SQLite Documentation[OL]. https://sqlite.org/docs.html, 2026.04.10",
    "[12] RAG_GPTV1.0 图扩展检索验收记录[Z]. 项目测试报告, 2026",
    "[13] RAG_GPTV1.0 重排模块对比记录[Z]. 项目测试报告, 2026",
]


def find_paragraph(document: Document, text: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def insert_paragraph_before(paragraph: Paragraph, text: str, style: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def delete_between(document: Document, start_text: str, end_text: str) -> None:
    start_index = next(i for i, p in enumerate(document.paragraphs) if p.text.strip() == start_text)
    end_index = next(i for i, p in enumerate(document.paragraphs) if p.text.strip() == end_text)
    for paragraph in list(document.paragraphs[start_index + 1 : end_index]):
        parent = paragraph._element.getparent()
        if parent is not None:
            parent.remove(paragraph._element)


def main() -> None:
    document = Document(DOC_PATH)
    delete_between(document, "参考文献", "致谢")
    end_paragraph = find_paragraph(document, "致谢")
    for item in REFERENCES:
        insert_paragraph_before(end_paragraph, item, "参考文献")
    temp_path = DOC_PATH.with_name(f"{DOC_PATH.stem}.refs.tmp.docx")
    if temp_path.exists():
        temp_path.unlink()
    document.save(temp_path)
    temp_path.replace(DOC_PATH)


if __name__ == "__main__":
    main()

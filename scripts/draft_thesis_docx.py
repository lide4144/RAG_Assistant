from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


@dataclass(frozen=True)
class Block:
    style: str
    text: str


TITLE = "面向论文知识库的规范驱动智能问答系统设计与实现"
ENGLISH_TITLE = "Design and Implementation of a Specification-Driven Intelligent Question Answering System for Paper Knowledge Bases"
OUTPUT_PATH = Path("output/doc/本科毕业论文初稿_20260409/附件6：大连海事大学本科毕业论文模板（理工类）-初稿.docx")

CHINESE_ABSTRACT = (
    "随着学术论文数量持续增长，围绕论文知识库的智能问答系统在检索准确性、证据可追溯性与多轮交互一致性方面面临更高要求。"
    "针对通用问答系统在论文场景中存在的证据混杂、引用不稳定以及证据不足时仍可能输出高风险结论等问题，"
    "本文结合已有项目代码、设计文档与阶段验收材料，设计并实现了一套面向论文知识库的规范驱动智能问答系统。"
    "系统采用 Frontend、Gateway 与 Python Kernel 三层架构，以 PDF 论文入库为起点，依次完成文本解析、Chunk 清洗、BM25 与向量索引构建、"
    "图扩展召回、候选重排、证据门控和带引用回答输出，并在运行态增加任务中心、配置治理和多轮会话能力。"
)

CHINESE_ABSTRACT_2 = (
    "在实现层面，系统前端基于 Next.js 构建聊天与任务工作台，网关使用 Express 与 WebSocket 统一编排流式事件，"
    "内核则基于 FastAPI 组织论文导入、索引构建、问答执行和运行日志落盘。根据现有阶段报告，图扩展模块在 10 个多跳与对比问题抽样中均补到了关键上下文；"
    "rerank 模块在 20 个问题对比中有 19 个问题的前 3 条证据更加相关，另 1 个至少保持持平。"
    "针对图扩展、rerank 与任务持久化的 17 项自动化测试全部通过。与此同时，多轮样式控制场景仍存在 2 个回归失败用例，"
    "说明对话控制与锚点复用在复杂问答链路中仍需继续优化。本文的工作表明，采用规范驱动和证据优先的工程路径，能够在本地论文知识库场景下构建兼顾可用性与可信性的问答系统。"
)

ENGLISH_ABSTRACT = (
    "As the number of academic papers keeps increasing, question answering systems for paper knowledge bases must provide better retrieval quality, traceable evidence, "
    "and stable multi-turn interaction. To address evidence mixing, unstable citation mapping, and risky conclusions under insufficient evidence, this thesis designs and "
    "implements a specification-driven intelligent question answering system for paper knowledge bases. The system adopts a three-layer architecture composed of a frontend, "
    "a gateway, and a Python kernel. Starting from PDF ingestion, it performs document parsing, chunk cleaning, BM25 and vector indexing, graph-based expansion, reranking, "
    "evidence gating, and citation-grounded answer generation."
)

ENGLISH_ABSTRACT_2 = (
    "The frontend is built with Next.js, the gateway uses Express and WebSocket to orchestrate streaming events, and the Python kernel exposes the retrieval and QA pipeline through FastAPI. "
    "Existing project reports show that graph expansion recovered key context for all sampled multi-hop and comparison questions, while reranking produced more relevant top-3 evidence on 19 out of 20 evaluated questions. "
    "In addition, 17 automated tests covering graph expansion, reranking, and job persistence passed successfully. However, two regression cases still fail in style-control follow-up scenarios, indicating that anchor reuse and dialogue control need further refinement. "
    "The implementation demonstrates that a specification-driven and evidence-first engineering route is effective for building a practical and trustworthy local paper QA system."
)

BODY_BLOCKS: list[Block] = [
    Block("Heading 1", "第1章 绪论"),
    Block("Heading 2", "1.1 课题背景与研究意义"),
    Block(
        "Normal",
        "随着大语言模型在知识问答、科研辅助和文档分析中的广泛应用，围绕专业文献构建可追溯问答系统已成为智能信息处理领域的重要工程方向。"
        "与面向开放互联网的通用问答不同，论文知识库场景更强调回答依据、引用位置和证据边界。用户不仅关注“系统能否回答”，还关注“答案来自哪一篇论文、哪一段文字以及是否超出了已有证据”。"
        "如果系统只输出流畅答案而无法给出稳定证据，就很难满足科研学习和论文阅读中的可信要求。"
    ),
    Block(
        "Normal",
        "本项目面向论文知识库场景，围绕“规范驱动、证据优先、渐进增强”的思路开展设计与实现。所谓规范驱动，是指需求、设计、任务和验收通过仓库中的规范文档持续管理；"
        "所谓证据优先，是指回答必须尽量绑定到检索到的文献片段，并在证据不足时主动降级而不是继续扩写。"
        "这一思路既有助于提升系统可信性，也有助于在毕业设计阶段形成可复现、可测试、可演进的工程闭环。"
    ),
    Block(
        "Normal",
        "从工程应用角度看，该系统能够为论文阅读、研究综述、对比分析和知识整理提供本地化工具支撑；从教学实践角度看，本项目展示了如何将检索增强生成、图扩展、会话状态管理、运行配置治理和多服务前后端协同整合为一个完整的软件系统。"
        "因此，开展本课题既具有较强的工程实践价值，也具有一定的研究与教学示范意义。"
    ),
    Block("Heading 2", "1.2 国内外研究现状"),
    Block(
        "Normal",
        "从已有公开技术路线看，论文问答系统的基础能力主要建立在检索增强生成框架之上。早期系统通常采用关键词检索或传统向量检索，将命中的文本片段直接交给生成模型总结。"
        "这类方法实现简单，但在多论文聚合、跨段落推理和引用约束方面存在明显不足，容易出现只命中局部结论、缺失实验条件或者引用编号与证据不一致等问题。"
    ),
    Block(
        "Normal",
        "为提升检索质量，当前常见做法是将 BM25、稠密向量检索以及重排模型组合使用，形成 Hybrid RAG 架构。"
        "其中，BM25 有利于命中关键术语和符号化表达，向量检索能够补足语义近邻，重排模块则进一步压缩候选集中的噪声。"
        "在此基础上，一些系统又引入图增强检索，通过邻接关系或实体关系扩展证据范围，以处理“方法与限制分别在哪里说明”“实验设置与结论是否对应”等多跳问题。"
    ),
    Block(
        "Normal",
        "另一方面，随着用户对可信问答的要求提高，仅依靠检索质量仍不足以解决风险问题。"
        "越来越多的系统开始关注证据门控、回答降级和多轮上下文管理，即在证据不足、问题指代不清或用户只提出格式控制要求时，系统应先识别交互意图，再决定是继续回答、请求澄清还是仅执行格式调整。"
        "本项目正是在这一背景下，将 Hybrid Retrieval、Graph Expansion、Evidence Policy、Sufficiency Gate 与多轮会话状态管理组合到统一工程中。"
    ),
    Block("Heading 2", "1.3 本文的主要研究内容"),
    Block(
        "Normal",
        "本文围绕一个面向论文知识库的规范驱动智能问答系统展开，主要研究内容包括以下几个方面。"
        "第一，构建从 PDF 入库、文本清洗、索引建立到问答输出的完整本地处理链路，形成稳定的数据准备基础。"
        "第二，设计融合 BM25、向量检索、图扩展和候选重排的检索侧方案，提高多论文、多跳问题下的证据覆盖能力。"
    ),
    Block(
        "Normal",
        "第三，围绕引用稳定性与回答安全性，设计证据分配、引用映射与 Sufficiency Gate 机制，使系统在生成答案时能够优先输出可追溯内容，并在证据不足时给出保守回应。"
        "第四，完成基于 Next.js、Node Gateway 和 Python Kernel 的三层产品化架构，实现流式回答、任务追踪、运行配置管理和开发者审查视图。"
    ),
    Block(
        "Normal",
        "全文结构安排如下：第 1 章说明课题背景、研究现状和本文工作；第 2 章介绍系统实现依赖的关键技术与理论基础；第 3 章分析系统需求；第 4 章给出总体架构与模块设计；"
        "第 5 章阐述关键模块实现过程；第 6 章结合自动化测试与阶段评估记录，对系统结果进行分析，并总结已知问题与后续改进方向。"
    ),
    Block("Heading 1", "第2章 相关技术与理论基础"),
    Block("Heading 2", "2.1 系统总体技术栈"),
    Block(
        "Normal",
        "本项目采用前端、网关和内核分层的技术组织方式。前端部分使用 Next.js 14 与 React 18 构建聊天工作台、任务中心与运行配置页面；"
        "网关层使用 Node.js、Express 与 WebSocket 负责浏览器连接管理、流式事件转发和本地/联网/混合三种问答模式的统一协议封装；"
        "内核部分使用 FastAPI 暴露问答接口、图构建任务接口和运行态管理接口，并调用已有的检索和生成模块完成核心计算。"
    ),
    Block(
        "Normal",
        "这种分层方式的优势在于职责明确。前端负责交互展示和状态管理，网关负责协议适配与实时传输，内核聚焦论文处理、检索推理和结果落盘。"
        "相比将所有能力堆叠在单体脚本中，三层结构更便于部署、调试和后续功能扩展，也更适合在毕业设计中展示一个完整软件系统的架构设计过程。"
    ),
    Block("Heading 2", "2.2 PDF 文档解析与文本清洗"),
    Block(
        "Normal",
        "论文知识库系统的前置条件是能够将 PDF 文档转化为结构化文本片段。项目中通过 `app.ingest`、`app.parser` 与 `app.marker_parser` 等模块完成论文入库处理，"
        "将原始论文转换为 `chunks.jsonl`、`papers.json` 等中间产物。随后，`app.clean_chunks` 对片段文本执行规范化清洗，包括空白处理、质量标注、类型识别和 `clean_text` 字段生成。"
    ),
    Block(
        "Normal",
        "清洗步骤的核心作用不在于简单去噪，而在于为后续检索建立统一语料基础。"
        "如果 BM25、TF-IDF 向量、Embedding 向量分别使用不同文本字段，系统就可能在召回时出现词项不一致与引用错位。"
        "因此，本项目把清洗后的 `clean_text` 作为重要中间表示，使索引构建、检索和证据输出尽量围绕同一文本源展开。"
    ),
    Block("Heading 2", "2.3 Hybrid RAG 检索理论"),
    Block(
        "Normal",
        "检索增强生成的基本思想是在生成答案之前先从外部知识源检索证据片段，再将这些片段送入模型进行归纳。"
        "针对论文场景，本项目未采用单一检索方式，而是同时构建 BM25 索引、TF-IDF 向量索引和可选 Embedding 向量索引。"
        "其中，BM25 对标题术语、方法名称和实验指标等关键词表达更敏感，向量检索则有利于召回语义相近但措辞不同的段落。"
    ),
    Block(
        "Normal",
        "Hybrid RAG 的关键不只是把多路结果简单拼接，而是让不同检索结果在统一候选空间中进行融合与去重。"
        "项目中的 `retrieve` 模块会综合 BM25、向量检索和可选 Embedding 路径形成候选集，再交由后续重排与证据门控处理。"
        "这一设计有利于平衡召回率与相关性，避免系统在面对中文口语化提问、多论文聚合和跨术语表达时只依赖单一路径。"
    ),
    Block("Heading 2", "2.4 图扩展与候选重排机制"),
    Block(
        "Normal",
        "论文中的关键信息往往分布在相邻段落、不同小节甚至参考部分中，仅靠初始检索结果可能无法覆盖完整证据链。"
        "因此，项目在初检之后引入 1-hop 图扩展能力，利用相邻关系和实体关系补全候选集合。"
        "这种做法更适合处理“定义在哪里，结果又在哪里”“限制说明是否对应结论”等需要跨段落关联的问题。"
    ),
    Block(
        "Normal",
        "为了避免扩展后候选规模失控，系统对新增候选设置了预算与过滤规则，并通过 rerank 模块进一步调整前列证据顺序。"
        "已有阶段报告显示，在 20 个问题的人工比对中，rerank 后的前 3 条证据有 19 个问题表现为更相关，剩余 1 个至少保持持平。"
        "说明图扩展与候选重排的组合能够在不显著放大噪声的情况下，提升最终输入到生成阶段的证据质量。"
    ),
    Block("Heading 2", "2.5 证据门控与多轮会话机制"),
    Block(
        "Normal",
        "论文问答系统不同于一般聊天机器人，回答不仅要自然，还要控制风险。"
        "本项目在回答生成前后加入 Evidence Policy 与 Sufficiency Gate 两类控制逻辑。前者约束关键结论必须与证据集合对应，后者在证据不足、问题范围不清或引用不稳定时触发保守策略，提示用户补充线索或仅输出可追溯部分。"
    ),
    Block(
        "Normal",
        "多轮会话方面，系统通过 `session_id`、历史摘要与锚点查询等机制，对“这篇论文”“他们”“用中文回答我”等后续输入进行上下文解释。"
        "这使得系统不仅能够处理事实问答，还能在一定程度上理解用户的格式控制、继续追问和研究辅助请求。"
        "不过，现有回归结果表明，样式控制与锚点复用在个别场景下仍存在失配，这也构成后续优化的重要方向。"
    ),
    Block("Heading 2", "2.6 本章小结"),
    Block(
        "Normal",
        "本章从技术基础层面对系统实现所依赖的核心方法进行了说明，包括三层服务架构、PDF 解析与文本清洗、Hybrid RAG 检索、图扩展与重排、证据门控以及多轮会话管理。"
        "这些技术共同构成了后续系统需求分析、架构设计和模块实现的基础。"
    ),
    Block("Heading 1", "第3章 需求分析"),
    Block("Heading 2", "3.1 系统建设目标"),
    Block(
        "Normal",
        "本系统的总体目标是构建一个面向论文知识库的本地智能问答平台，使用户能够围绕已导入论文提出事实性问题、对比性问题与开放式研究辅助问题，并获得带引用、可复查的回答。"
        "与面向开放网络的一般搜索产品相比，本系统更关注论文证据约束和回答边界控制，即系统应优先保证结果可信，再逐步提升交互体验与功能丰富度。"
    ),
    Block(
        "Normal",
        "围绕上述目标，项目需要同时满足算法链路完整和产品交互可用两类要求。"
        "前者要求系统完成文档入库、索引建立、检索、重排、生成和日志记录；后者要求系统具备聊天界面、任务反馈、引用展示、运行状态概览与配置管理等能力。"
        "只有将两类目标结合，系统才具备面向真实使用场景的可落地性。"
    ),
    Block("Heading 2", "3.2 功能需求分析"),
    Block(
        "Normal",
        "结合现有仓库功能，系统的核心功能需求可以归纳为五类。第一类是知识库管理需求，包括论文 PDF 导入、入库结果反馈、论文目录维护和中间产物生成。"
        "第二类是检索与问答需求，包括本地问答、联网补充问答、混合问答、证据分组展示和引用编号映射。"
        "第三类是推理增强需求，包括 Query Rewriting、意图校准、图扩展、候选重排和证据门控。"
    ),
    Block(
        "Normal",
        "第四类是运行控制需求，包括图构建任务提交、后台任务查询、运行配置保存、模型配置检测和失败诊断。"
        "第五类是会话交互需求，包括多轮历史保持、上下文清空、开发者审查视图和流式回答体验。"
        "这些功能共同支撑用户从导入论文、建立索引到发起问答和分析结果的完整使用流程。"
    ),
    Block("Heading 2", "3.3 非功能需求分析"),
    Block(
        "Normal",
        "论文知识库问答系统除了功能正确外，还需要满足一定的非功能要求。首先是可追溯性，系统输出的关键结论应尽可能附带来源信息，并保证引用编号与证据列表保持一致。"
        "其次是可维护性，系统的配置、日志、测试和规范文档需要彼此对应，便于后续迭代与问题排查。"
    ),
    Block(
        "Normal",
        "再次是可扩展性。由于项目后期逐步引入 Planner、任务中心、联网路径和可选 Marker 解析，因此架构设计必须允许新增模块在不破坏主链路的前提下接入。"
        "最后是交互可理解性，系统应通过状态提示、流式输出、错误消息和审查信息帮助用户理解当前执行过程，避免出现“看起来能回答，实际上证据不足”的误用。"
    ),
    Block("Heading 2", "3.4 业务流程分析"),
    Block(
        "Normal",
        "从用户视角看，系统的典型业务流程包括：上传或导入论文，系统解析 PDF 并生成清洗后的文本片段；随后构建检索索引和可选图结构；用户在聊天界面选择本地、联网或混合模式并提出问题；"
        "系统根据历史上下文分析问题意图，选择合适的检索与回答路径；最后将答案、引用、运行日志和可视化调试信息返回给前端。"
    ),
    Block(
        "Normal",
        "从内部执行角度看，问答流程可概括为“问题输入、意图分析、检索召回、图扩展、重排、证据门控、答案生成、引用映射、结果落盘”九个阶段。"
        "这一流程说明系统不仅是一个简单的聊天页面，而是一个围绕文献证据组织起来的多阶段处理系统。"
    ),
    Block("Heading 2", "3.5 本章小结"),
    Block(
        "Normal",
        "本章从系统建设目标、功能需求、非功能需求和业务流程四个角度对项目需求进行了分析。"
        "这些分析表明，系统既要满足论文知识库场景下的检索与问答能力，也要具备产品化部署、日志治理与可维护性要求，为后续架构设计提供了依据。"
    ),
    Block("Heading 1", "第4章 系统设计"),
    Block("Heading 2", "4.1 总体架构设计"),
    Block(
        "Normal",
        "系统总体采用三层架构。最上层为前端展示层，承担聊天交互、任务面板、运行态概览和设置页等界面功能；中间层为网关编排层，负责 HTTP 与 WebSocket 协议接入、事件流转发和模式分流；"
        "底层为 Python Kernel 能力层，负责文档入库、索引构建、问答执行、任务持久化和配置读写。"
    ),
    Block(
        "Normal",
        "这种设计将“用户交互逻辑”和“文档问答能力”解耦开来。前端无需直接理解复杂检索链路，只需围绕统一事件协议渲染消息和引用；"
        "网关也不直接参与论文检索，而是负责将本地 SSE 结果转化为浏览器可消费的 WebSocket 消息；内核则通过独立接口输出问答和后台任务结果。"
        "在工程实践中，这种分层更有利于单独测试和定位问题。"
    ),
    Block("Heading 2", "4.2 数据流与处理流程设计"),
    Block(
        "Normal",
        "系统数据流可以分为离线准备流和在线问答流两部分。离线准备流从 `data/papers` 中读取 PDF 文件，经 `ingest` 生成 `chunks.jsonl` 与 `papers.json`，再经 `clean_chunks` 产生 `chunks_clean.jsonl`，"
        "随后建立 `bm25_index.json`、`vec_index.json` 和可选的 `vec_index_embed.json`。若启用图构建，还会生成 `graph.json` 作为图扩展阶段的邻接来源。"
    ),
    Block(
        "Normal",
        "在线问答流以用户输入为起点，首先完成历史上下文合并、意图识别和查询改写；然后载入索引进行候选召回，并根据配置执行图扩展与 rerank；"
        "完成证据分配与充足性判断后，系统再调用生成模块输出答案，同时写入 `run_trace.json` 与 `qa_report.json` 等运行日志。"
        "这一设计保证了每轮问答都能留存可复查痕迹。"
    ),
    Block("Heading 2", "4.3 核心模块设计"),
    Block(
        "Normal",
        "内核侧核心模块主要包括文档入库模块、索引构建模块、检索与重排模块、问答生成模块、会话状态模块和任务持久化模块。"
        "其中，`app.ingest`、`app.clean_chunks`、`app.build_indexes` 负责离线数据准备；`app.retrieve`、`app.rerank`、`app.qa` 负责在线问答主链路；"
        "`app.session_state` 用于管理多轮上下文；`app.job_store` 则使用 SQLite 记录后台任务状态与事件。"
    ),
    Block(
        "Normal",
        "网关侧模块主要包括协议解析、流式消息转发、引用稳定化处理和联网检索适配。前端侧则以聊天壳层、图子图面板、任务中心、Pipeline Workbench 和设置页面为主要组成。"
        "从模块边界上看，项目较好地区分了“算法模块”“接口模块”和“交互模块”，有助于后续局部替换和独立演进。"
    ),
    Block("Heading 2", "4.4 数据与存储设计"),
    Block(
        "Normal",
        "系统没有采用传统关系数据库来承载全部论文内容，而是根据不同数据类型选择更轻量的文件化存储方式。"
        "论文目录和处理中间结果以 JSON 或 JSONL 形式存储，既方便调试也便于脚本化处理；检索索引与图结构分别落盘为独立文件，避免每轮问答重复构建。"
        "这一设计符合毕业设计阶段“小规模、本地化、强调可追溯”的工程特点。"
    ),
    Block(
        "Normal",
        "对于需要持久化状态的后台任务，系统使用 `job_store.sqlite3` 存储任务记录、事件序列和配置快照。"
        "SQLite 的引入使任务中心能够在服务重启后继续保留必要信息，同时不必额外部署独立数据库服务。"
        "从实现复杂度、部署成本与功能需求的平衡来看，这一方案较为合适。"
    ),
    Block("Heading 2", "4.5 接口与运行配置设计"),
    Block(
        "Normal",
        "系统通过 FastAPI 暴露问答、图构建、任务查询和运行配置等接口，并由网关进一步封装成浏览器使用的统一入口。"
        "对于聊天主链路，内核支持 `/qa` 与 `/qa/stream` 两种接口形式，分别用于一次性结果返回和基于 SSE 的流式响应。"
        "网关接收这些结果后，将其转换为 `message`、`sources`、`messageEnd` 和 `error` 等 WebSocket 事件。"
    ),
    Block(
        "Normal",
        "配置方面，系统区分默认配置、运行态配置和环境变量覆盖三层来源。"
        "这种治理方式使模型 API、图构建参数和 Marker 开关等配置既可在本地文件中统一维护，也可在运行中通过管理接口动态调整。"
        "配置来源的显式区分，对于保证实验可复现和避免部署时配置漂移具有重要作用。"
    ),
    Block("Heading 2", "4.6 本章小结"),
    Block(
        "Normal",
        "本章从总体架构、数据流、核心模块、数据存储和接口配置五个方面给出了系统设计。"
        "通过三层架构、文件化中间产物与 SQLite 任务持久化的组合，系统在保证开发效率的同时，为后续实现与测试提供了清晰的结构基础。"
    ),
    Block("Heading 1", "第5章 系统实现"),
    Block("Heading 2", "5.1 开发与运行环境"),
    Block(
        "Normal",
        "项目主要使用 Python、TypeScript 和 Node.js 三类运行环境。Python 侧依赖 FastAPI、PyYAML、LiteLLM、PyMuPDF 等库，负责问答链路、配置解析和 PDF 处理；"
        "前端侧使用 Next.js 14、React 18、Tailwind CSS 等技术完成页面渲染；网关侧基于 Express、Axios 和 WebSocket 完成服务编排。"
        "从仓库结构看，前端、网关和内核依赖分别隔离，便于独立安装与启动。"
    ),
    Block(
        "Normal",
        "在本地开发模式下，三个服务分别监听 3000、8080 和 8000 端口，并通过脚本统一启动。"
        "这一端口约定既方便浏览器开发调试，也有利于在论文中说明系统的分层部署关系。"
        "当需要更严格的联网行为时，还可以通过环境变量显式开启真实 Web Provider 与严格失败模式。"
    ),
    Block("Heading 2", "5.2 文档入库与索引实现"),
    Block(
        "Normal",
        "在文档入库实现中，系统首先扫描 `data/papers` 目录中的 PDF 文件，并调用解析模块生成论文元数据与分块文本。"
        "每个文本片段在后续清洗过程中会补充内容类型、质量标记和标准化文本字段。"
        "这一实现不仅服务于后续检索，也为引用输出提供了稳定的片段标识。"
    ),
    Block(
        "Normal",
        "索引构建阶段由 `app.build_indexes` 统一驱动，可同时生成 BM25 索引、TF-IDF 向量索引和可选的 Embedding 索引。"
        "若启用图构建，`app.graph_build` 还会根据片段邻接和实体关系形成图结构。"
        "这一实现思路强调离线预处理，避免每次问答都重新读取全文和实时建模，从而提高系统的交互响应效率。"
    ),
    Block("Heading 2", "5.3 问答主链路实现"),
    Block(
        "Normal",
        "问答主链路由 `app.qa` 组织。该模块会先执行问题标准化、查询改写与意图校准，再载入检索索引完成候选召回。"
        "若配置允许，系统会继续执行图扩展与候选重排，之后进入证据分配和 Sufficiency Gate 判断阶段。"
        "只有在关键结论具备足够证据支撑时，生成模块才输出完整回答；否则系统将返回保守提示或建议补充线索。"
    ),
    Block(
        "Normal",
        "从实现角度看，`app.qa` 并不是简单串联多个函数，而是承担了多个阶段状态的协调者角色。"
        "它不仅负责调度检索与生成，还要记录运行目录、输出调试字段、维护引用映射和整合多轮上下文。"
        "因此，该模块是整个项目中最能体现系统复杂度和工程组织能力的核心组件。"
    ),
    Block("Heading 2", "5.4 多服务交互实现"),
    Block(
        "Normal",
        "前端聊天页面通过统一消息流与网关交互，用户可在“仅知识库”“联网补充”“混合回答”三种模式间切换。"
        "当用户发起问题后，网关会根据模式选择直接调用本地内核、外部联网检索或二者混合策略，并将响应拆解为适合前端逐步渲染的事件。"
        "这种事件流式设计使用户能够在答案尚未完全生成时看到中间输出，提升等待期间的可理解性。"
    ),
    Block(
        "Normal",
        "网关中的 `chatService` 除了转发消息外，还承担引用顺序稳定化、错误事件封装和后台任务轮询等职责。"
        "前端组件则进一步将回答、引用、图子图、调试信息和任务状态组织为统一界面。"
        "从实现效果看，多服务交互层把复杂的检索问答链路转化为了可消费的用户体验，是系统产品化的重要一步。"
    ),
    Block("Heading 2", "5.5 运行配置与任务中心实现"),
    Block(
        "Normal",
        "为支持长耗时任务与运行态治理，系统在内核中实现了任务持久化机制。"
        "图构建、批量导入等操作可以先创建后台任务，再由前端任务中心持续查询状态和事件序列。"
        "SQLite 任务库保存了任务基本信息、更新时间、进度阶段、输出摘要和配置快照，从而避免浏览器刷新后完全失去任务上下文。"
    ),
    Block(
        "Normal",
        "此外，系统还提供运行态配置管理接口，用于保存和读取模型 API、Planner 配置、Marker 配置与启停开关。"
        "这些实现一方面增强了项目的可维护性，另一方面也为后续论文中的实验复现实验和部署说明提供了基础。"
        "相较于只依靠环境变量临时调整参数的方式，运行配置中心更适合持续开发阶段的系统管理。"
    ),
    Block("Heading 2", "5.6 本章小结"),
    Block(
        "Normal",
        "本章围绕开发环境、文档入库与索引、问答主链路、多服务交互以及运行配置与任务中心等内容说明了系统的关键实现。"
        "通过这些实现，项目已经具备从论文导入、知识组织到带引用问答输出的完整能力，也为系统测试与结果分析奠定了基础。"
    ),
    Block("Heading 1", "第6章 系统测试与结果分析"),
    Block("Heading 2", "6.1 测试环境与测试方法"),
    Block(
        "Normal",
        "本项目的测试主要由三部分组成。第一部分是单元与回归测试，用于验证图扩展、重排、任务持久化、多轮会话和接口契约等关键模块。"
        "第二部分是阶段评估报告，围绕图扩展召回效果、rerank 相关性和研究助手工作流进行人工抽样验证。"
        "第三部分是人工验收清单，用于约束导入论文、发起问答、查看引用和保存草稿等端到端交互行为。"
    ),
    Block(
        "Normal",
        "从测试环境上看，项目在本地开发环境中运行三层服务，并通过 `pytest` 执行 Python 侧自动化测试。"
        "论文写作阶段，本文额外在当前仓库中执行了 `tests/test_m5_graph_expansion.py`、`tests/test_rerank.py` 和 `tests/test_job_store.py` 三组关键测试，用于验证图扩展、重排和任务持久化路径。"
    ),
    Block("Heading 2", "6.2 核心功能验证结果"),
    Block(
        "Normal",
        "在自动化测试方面，上述三组关键测试共 17 个用例全部通过，执行时间约为 2.32 秒。"
        "这说明图扩展候选补全、rerank 结果组织和 SQLite 任务持久化模块在当前版本下处于稳定可用状态。"
        "同时，测试输出中出现 3 条与 FastAPI 生命周期写法和 `python_multipart` 导入方式有关的警告，表明部分基础依赖和框架用法仍有清理空间。"
    ),
    Block(
        "Normal",
        "在图扩展阶段评估中，现有人工抽样报告针对 10 个多跳与对比问题比较了扩展前后的证据覆盖情况。"
        "结果显示，扩展后 10 个样例全部补到了关键上下文，包括局限段、实验设置段、指标定义段和参考部分等。"
        "说明 1-hop 图扩展能够有效弥补初检候选只命中局部段落的问题。"
    ),
    Block(
        "Normal",
        "在 rerank 阶段评估中，项目以 20 个问题作为人工对比样例，观察 rerank 前后前 3 条证据是否更贴合问题语义。"
        "结果表明，20 个问题中有 19 个表现为“更相关”，剩余 1 个至少与调整前持平。"
        "这一结果说明，经过图扩展后的候选集仍需要重排处理，而重排可以进一步抑制噪声候选进入生成阶段。"
    ),
    Block("Heading 2", "6.3 结果分析"),
    Block(
        "Normal",
        "综合现有结果可以看出，系统在“证据覆盖”和“证据排序”两个关键维度上取得了较明显改进。"
        "图扩展补齐了跨段落、多跳和参考信息相关问题的上下文缺口，而 rerank 则改善了最终提交给生成模块的候选质量。"
        "这两个环节共同作用，使系统更适合处理论文问答中常见的定义、实验、局限和对比类问题。"
    ),
    Block(
        "Normal",
        "另一方面，项目在产品化层面也形成了较为完整的支撑体系。"
        "前端工作台、网关流式事件、任务中心和配置治理共同提升了系统的可观测性，使开发者能够从运行态概览、调试记录和任务事件中定位问题。"
        "对于毕业设计而言，这一点非常重要，因为它体现的不是单一算法点，而是围绕实际系统可维护性展开的工程能力。"
    ),
    Block("Heading 2", "6.4 已知问题与改进方向"),
    Block(
        "Normal",
        "在补充回归测试时，本文将 `tests/test_m7_6_multi_turn.py` 与其它核心测试一并执行，结果出现 2 个失败用例。"
        "失败场景均与样式控制型追问有关，具体表现为在证据不足提示语出现后，系统对最近主题锚点与引用片段的复用不符合预期。"
        "这说明多轮对话中的“内容问题”和“格式控制问题”在当前实现中仍存在耦合。"
    ),
    Block(
        "Normal",
        "后续改进可从三个方面展开。第一，进一步分离样式控制与事实问答的状态更新逻辑，避免格式指令污染上一轮主题锚点。"
        "第二，在 Sufficiency Gate 触发保守回答时，为会话状态保存更明确的证据锚点，减少后续跟进中出现的歧义。"
        "第三，继续补充端到端自动化测试，将前端交互、网关事件和内核状态统一纳入回归范围，以更早发现跨层问题。"
    ),
    Block("Heading 2", "6.5 本章小结"),
    Block(
        "Normal",
        "本章结合自动化测试与阶段评估记录，对系统的核心结果进行了分析。"
        "现有证据表明，图扩展、rerank 和任务持久化等关键能力已经具备较好的稳定性；与此同时，多轮样式控制场景仍暴露出可继续完善的问题。"
        "因此，系统总体上达到了毕业设计初稿阶段的实现目标，但仍需在对话控制与端到端回归方面继续打磨。"
    ),
    Block("Heading 1", "结论与展望"),
    Block(
        "Normal",
        "本文围绕面向论文知识库的规范驱动智能问答系统，完成了从需求分析、架构设计到关键模块实现与测试分析的初步整理。"
        "项目已经实现论文入库、文本清洗、Hybrid RAG 检索、图扩展、候选重排、证据门控、三层服务协同和任务中心等能力，形成了一个较完整的本地论文问答系统原型。"
        "结合现有阶段报告与自动化测试结果可以看出，系统在证据覆盖、证据排序和运行可观测性方面取得了较好的工程效果。"
    ),
    Block(
        "Normal",
        "与此同时，本文也发现系统在多轮样式控制、复杂上下文锚点复用和端到端跨层验证方面仍有不足。"
        "后续工作将围绕三个方向继续展开：一是补强会话控制和 Planner 相关链路的稳定性；二是完善界面层与调试层的联动展示；三是结合最终答辩前的实验要求，补充更规范的人工验收记录、插图和参考文献格式。"
        "通过这些工作，系统有望进一步提升成熟度，并为论文终稿提供更完整的材料支撑。"
    ),
    Block("Heading 1", "参考文献"),
    Block("Normal", "以下参考文献为论文初稿阶段整理条目，后续需结合导师意见进一步核对页码、卷期和学校格式要求。"),
    Block("参考文献内容", "[1] Lewis P, Perez E, Piktus A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks[C]. Advances in Neural Information Processing Systems, 2020."),
    Block("参考文献内容", "[2] Robertson S, Zaragoza H. The Probabilistic Relevance Framework: BM25 and Beyond[J]. Foundations and Trends in Information Retrieval, 2009, 3(4):333-389."),
    Block("参考文献内容", "[3] Vaswani A, Shazeer N, Parmar N, et al. Attention Is All You Need[C]. Advances in Neural Information Processing Systems, 2017."),
    Block("参考文献内容", "[4] RAG_GPTV1.0 项目 README[Z]. /home/programer/RAG_GPTV1.0/README.md."),
    Block("参考文献内容", "[5] 本科毕业设计（论文）开题报告[Z]. /home/programer/RAG_GPTV1.0/docs/本科毕业设计开题报告.md."),
    Block("参考文献内容", "[6] Multi-Service Local Development[Z]. /home/programer/RAG_GPTV1.0/docs/multi-service-dev.md."),
    Block("参考文献内容", "[7] M5 Graph Expansion Retrieval 验收记录[Z]. /home/programer/RAG_GPTV1.0/reports/m5_graph_expansion_eval.md."),
    Block("参考文献内容", "[8] M6 Rerank 对比记录[Z]. /home/programer/RAG_GPTV1.0/reports/m6_rerank.md."),
    Block("参考文献内容", "[9] Research Assistant Manual Acceptance[Z]. /home/programer/RAG_GPTV1.0/reports/research_assistant_manual_acceptance.md."),
    Block("参考文献内容", "[10] FastAPI Documentation[EB/OL]. https://fastapi.tiangolo.com/."),
    Block("Heading 1", "致谢"),
    Block(
        "Normal",
        "本论文初稿的完成离不开项目现有代码、设计文档、阶段报告和测试材料所提供的基础。"
        "在后续正式定稿阶段，致谢部分还需结合指导教师意见、课程支持情况以及个人实践过程进一步补充和润色。"
    ),
    Block("Heading 1", "附录"),
    Block("Heading 2", "附录1：主要运行产物说明"),
    Block(
        "Normal",
        "为便于后续完善论文，现将本项目在撰写阶段涉及的主要运行产物归纳如下：`data/processed/chunks.jsonl` 和 `data/processed/chunks_clean.jsonl` 用于保存论文片段及清洗结果；"
        "`data/indexes/bm25_index.json`、`data/indexes/vec_index.json` 与 `data/indexes/vec_index_embed.json` 用于保存多路检索索引；"
        "`data/processed/graph.json` 用于保存图扩展结构；`runs/*` 目录用于记录问答过程中的 `run_trace.json` 与 `qa_report.json`。"
    ),
    Block(
        "Normal",
        "此外，`data/processed/job_store.sqlite3` 用于记录任务中心所需的任务状态与事件序列，`configs/llm_runtime_config.json` 与 `configs/planner_runtime_config.json` 用于保存运行态配置。"
        "这些文件路径既是系统实现的重要组成部分，也可在论文终稿中作为实现说明和复现实验的支撑材料。"
    ),
]


def find_paragraph(document: Document, text: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def insert_paragraph_before(paragraph: Paragraph, text: str, style: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def clear_runs(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink") or child.tag.endswith("}fldSimple"):
            paragraph._p.remove(child)


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    clear_runs(paragraph)
    paragraph.add_run(text)


def delete_paragraph(paragraph: Paragraph) -> None:
    parent = paragraph._element.getparent()
    if parent is not None:
        parent.remove(paragraph._element)


def delete_between(document: Document, start_text: str, end_text: str, *, inclusive: bool = False) -> None:
    start_index = next(i for i, p in enumerate(document.paragraphs) if p.text.strip() == start_text)
    end_index = next(i for i, p in enumerate(document.paragraphs) if p.text.strip() == end_text)
    paragraphs = document.paragraphs[start_index : end_index + 1]
    if not inclusive:
        paragraphs = paragraphs[1:-1]
    for paragraph in list(paragraphs):
        delete_paragraph(paragraph)


def replace_section_content(document: Document, start_text: str, end_text: str, blocks: list[Block]) -> None:
    delete_between(document, start_text, end_text, inclusive=False)
    end_paragraph = find_paragraph(document, end_text)
    for block in blocks:
        insert_paragraph_before(end_paragraph, block.text, block.style)


def remove_after_paragraph(paragraph: Paragraph) -> None:
    body = paragraph._element.getparent()
    if body is None:
        return
    seen = False
    for child in list(body):
        if child == paragraph._element:
            seen = True
            continue
        if not seen:
            continue
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def set_cell_text(cell, text: str) -> None:
    cell.text = text
    if cell.paragraphs:
        cell.paragraphs[0].style = "Normal"


def ensure_update_fields(document: Document) -> None:
    settings = document.settings.element
    if settings.find(qn("w:updateFields")) is not None:
        return
    element = OxmlElement("w:updateFields")
    element.set(qn("w:val"), "true")
    settings.append(element)


def fill_cover(document: Document) -> None:
    cover_title = document.tables[0]
    set_cell_text(cover_title.cell(0, 1), TITLE)
    set_cell_text(cover_title.cell(1, 1), ENGLISH_TITLE)

    info_table = document.tables[1]
    values = [
        "待填写",
        "待填写",
        "信息科学技术学院",
        "计算机科学与技术专业 2022 级 待填写班",
        "待填写（职称）",
        "待填写（职称）",
        "2026年4月",
    ]
    for row, value in zip(info_table.rows, values, strict=False):
        set_cell_text(row.cells[1], value)


def fill_headers(document: Document) -> None:
    document.settings.odd_and_even_pages_header_footer = True
    for section in document.sections:
        if section.header.paragraphs:
            set_paragraph_text(section.header.paragraphs[0], "大连海事大学本科毕业论文")
        if section.even_page_header.paragraphs:
            set_paragraph_text(section.even_page_header.paragraphs[0], TITLE)


def build_body(document: Document) -> None:
    anchor = find_paragraph(document, "大连海事大学本科毕业论文模板（理工类）")
    set_paragraph_text(anchor, TITLE)
    remove_after_paragraph(anchor)
    for block in BODY_BLOCKS:
        paragraph = document.add_paragraph(style=block.style)
        paragraph.add_run(block.text)


def clean_front_matter(document: Document) -> None:
    statement_index = next(i for i, p in enumerate(document.paragraphs) if p.text.strip() == "毕业论文原创性声明")
    following = document.paragraphs[statement_index + 1]
    set_paragraph_text(
        following,
        f"本人郑重声明：所呈交的毕业论文《{TITLE}》，是在导师指导下独立进行研究工作所取得的成果。除文中已经注明引用的内容外，"
        "本论文不包含其他个人或集体已经发表或撰写过的研究成果。对本文研究做出重要贡献的个人和集体，均已在文中以明确方式标明。"
        "本人对本声明的法律责任承担相应后果。"
    )
    delete_between(document, "明确以下说明后，删除此处红色文字。", "明确以上说明后，删除此处红色文字。", inclusive=True)


def main() -> None:
    document = Document(OUTPUT_PATH)
    document.core_properties.title = TITLE
    document.core_properties.subject = "本科毕业论文初稿"
    document.core_properties.keywords = "论文知识库, 智能问答, Hybrid RAG, 图扩展, 证据门控"

    fill_cover(document)
    clean_front_matter(document)
    replace_section_content(
        document,
        "摘    要",
        "ABSTRACT",
        [
            Block("Normal", CHINESE_ABSTRACT),
            Block("Normal", CHINESE_ABSTRACT_2),
            Block("Normal", "关键词：论文知识库；智能问答；Hybrid RAG；图扩展；证据门控"),
        ],
    )
    replace_section_content(
        document,
        "ABSTRACT",
        "目    录",
        [
            Block("Normal", ENGLISH_ABSTRACT),
            Block("Normal", ENGLISH_ABSTRACT_2),
            Block("Normal", "Keywords: paper knowledge base; intelligent question answering; hybrid RAG; graph expansion; evidence gating"),
        ],
    )
    build_body(document)
    fill_headers(document)
    ensure_update_fields(document)
    temp_path = OUTPUT_PATH.with_name(f"{OUTPUT_PATH.stem}.tmp.docx")
    if temp_path.exists():
        temp_path.unlink()
    document.save(temp_path)
    temp_path.replace(OUTPUT_PATH)


if __name__ == "__main__":
    main()
